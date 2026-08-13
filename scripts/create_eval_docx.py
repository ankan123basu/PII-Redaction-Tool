import re
from docx import Document
from docx.shared import Pt, Inches

def md_to_docx(md_path, docx_path):
    doc = Document()
    doc.add_heading("Evaluation Strategy & Metrics", 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("|") and "---" not in line:
            # Very basic table line handling (just add as text for simplicity to ensure it generates)
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = clean_line.replace("|", " ").strip()
            doc.add_paragraph(clean_line)
        elif not line.startswith("|"):
            # Strip bold formatting for standard text
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            doc.add_paragraph(clean_line)
            
    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

if __name__ == "__main__":
    md_to_docx("EVALUATION_REPORT.md", "data/output/Evaluation_Strategy_and_Metrics.docx")
