"""
Generate a synthetic sample_input.docx fixture for testing.

Creates a DOCX with known PII in paragraphs, tables, headers, and footers
so tests can verify detection and redaction against known ground truth.

Run this script to regenerate the fixture:
    python tests/fixtures/generate_fixture.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_sample_docx(output_path: str | Path = None) -> Path:
    """Generate a sample DOCX with known PII for testing."""
    if output_path is None:
        output_path = Path(__file__).parent / "sample_input.docx"

    doc = Document()

    # --- Section header/footer ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.text = "TechStart Innovations Pvt. Ltd. | CIN: U28129PN1979PLC141032"

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.text = "Confidential — Contact: info@techstartup.co.in | +91 98765 43210"

    # --- Cover page ---
    title = doc.add_heading("RED HERRING PROSPECTUS", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "TechStart Innovations Pvt. Ltd. (CIN: U28129PN1979PLC141032) "
        "is a company incorporated under the Companies Act, 2013. "
        "The registered office of the Company is located at "
        "42, MG Road, Bangalore, Karnataka 560001. "
        "For inquiries, contact info@techstartup.co.in or call +91 98765 43210."
    )

    # --- Promoters section ---
    doc.add_heading("Promoters and Promoter Group", level=2)
    doc.add_paragraph(
        "The promoters of the Company are Mr. Kushal Subbayya Hegde "
        "(PAN: ABCPD1234E, Phone: +91-9876543210) and "
        "Ms. Rashi Patil. Mr. Vijay Kumar serves as an independent director. "
        "Kushal Hegde has been associated with the Company since its incorporation. "
        "Date of Birth (DOB): 15/03/1985."
    )

    # --- Directors table ---
    doc.add_heading("Board of Directors", level=2)
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"

    # Header row
    headers = ["Name", "Designation", "PAN", "Contact", "Address"]
    for i, header_text in enumerate(headers):
        table.rows[0].cells[i].text = header_text

    # Director 1
    row1 = table.rows[1].cells
    row1[0].text = "Kushal Subbayya Hegde"
    row1[1].text = "Managing Director"
    row1[2].text = "ABCPD1234E"
    row1[3].text = "+91 98765 43210"
    row1[4].text = "42, MG Road\nBangalore, Karnataka\n560001"

    # Director 2
    row2 = table.rows[2].cells
    row2[0].text = "Dr. Anita Sharma"
    row2[1].text = "Independent Director"
    row2[2].text = "BCRPH5678F"
    row2[3].text = "+91 87654 32109"
    row2[4].text = "Plot No. 15, MIDC Industrial Area\nPune, Maharashtra\n411001"

    # Director 3
    row3 = table.rows[3].cells
    row3[0].text = "Rashi Patil"
    row3[1].text = "Whole-time Director"
    row3[2].text = "DEFGH9012I"
    row3[3].text = "+91 76543 21098"
    row3[4].text = "18, Park Avenue\nMumbai, Maharashtra\n400001"

    # --- Subsidiary section ---
    doc.add_heading("Subsidiary Companies", level=2)
    doc.add_paragraph(
        "The Company's wholly owned subsidiary, Tech Solutions Corp. "
        "(CIN: L67890MH2005PLC123456), is registered in Mumbai."
    )

    # --- Financial section with non-PII numbers ---
    doc.add_heading("Financial Information", level=2)
    doc.add_paragraph(
        "The Company was incorporated on January 15, 2005 under the "
        "Companies Act, 1956. The Company's authorized share capital is "
        "INR 50,00,00,000 (Rupees Fifty Crores). The filing was made on "
        "March 22, 2024. The ISIN for the equity shares is INE123456789. "
        "Application No. 2024-IPO-0042."
    )

    # --- Registrar section ---
    doc.add_heading("Registrar and Transfer Agent", level=2)
    doc.add_paragraph(
        "Global Registrars Limited\n"
        "Email: compliance@registrar.com\n"
        "The Board of Directors has approved the appointment of the Registrar."
    )

    # --- Advisors section ---
    doc.add_heading("Book Running Lead Managers", level=2)
    doc.add_paragraph(
        "SBI Capital Markets Ltd. has been appointed as the lead manager. "
        "The Company has engaged multiple advisors for the IPO process."
    )

    # --- Tech / IP section ---
    doc.add_heading("Technology Infrastructure", level=2)
    doc.add_paragraph(
        "The Company operates its primary servers at IP address 192.168.1.100. "
        "The system runs version 2.4.1 of the proprietary software."
    )

    # --- SSN and CC sample (rare in Indian docs but needed for completeness) ---
    doc.add_heading("International Operations", level=2)
    doc.add_paragraph(
        "For US operations, the EIN is 123-45-6789. "
        "Payment processing test card: 4111 1111 1111 1111."
    )

    doc.save(str(output_path))
    return Path(output_path)


if __name__ == "__main__":
    path = generate_sample_docx()
    print(f"Generated: {path}")
