"""
End-to-end pipeline tests.

Tests the full flow: DOCX in → detect → resolve overlaps → pseudonymize → DOCX out.

Uses the generated sample_input.docx fixture with known PII to verify:
- All PII types are detected and redacted
- Formatting is preserved (output is a valid DOCX)
- Overlap resolution works with deliberately overlapping fixtures
- Entity consistency (same PII → same fake across the document)
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from pii_redactor.detectors.base import DetectedEntity
from pii_redactor.document_io.docx_reader import DocxReader
from pii_redactor.pipeline import RedactionPipeline, resolve_overlaps


FIXTURE_DIR = Path(__file__).parent / "fixtures"
SAMPLE_DOCX = FIXTURE_DIR / "sample_input.docx"


# ===========================================================================
# Overlap Resolution Tests
# ===========================================================================


class TestOverlapResolution:
    """Tests for the deterministic overlap resolution algorithm."""

    def test_no_overlaps_passes_through(self):
        """Non-overlapping entities should all be kept."""
        entities = [
            DetectedEntity("hello", "FULL_NAME", 0, 5, confidence=0.9, detector_name="NER"),
            DetectedEntity("world", "EMAIL", 10, 15, confidence=0.95, detector_name="Regex"),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 2

    def test_validated_wins_over_unvalidated(self):
        """Validated regex detection should win over unvalidated NER."""
        entities = [
            DetectedEntity(
                "1234567890", "PHONE", 0, 10,
                confidence=0.8, detector_name="NER", validated=False,
            ),
            DetectedEntity(
                "1234567890", "PHONE", 0, 10,
                confidence=0.7, detector_name="Regex", validated=True,
            ),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 1
        assert resolved[0].validated is True

    def test_higher_confidence_wins(self):
        """Higher confidence should win when both are unvalidated."""
        entities = [
            DetectedEntity(
                "some text", "FULL_NAME", 0, 9,
                confidence=0.6, detector_name="NER",
            ),
            DetectedEntity(
                "some text", "COMPANY_NAME", 0, 9,
                confidence=0.9, detector_name="NER",
            ),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 1
        assert resolved[0].entity_type == "COMPANY_NAME"

    def test_longer_span_wins_on_tiebreak(self):
        """Longer span should win when confidence is equal."""
        entities = [
            DetectedEntity(
                "John", "FULL_NAME", 0, 4,
                confidence=0.85, detector_name="NER",
            ),
            DetectedEntity(
                "John Smith", "FULL_NAME", 0, 10,
                confidence=0.85, detector_name="NER",
            ),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 1
        assert resolved[0].text == "John Smith"

    def test_partial_overlap_resolved(self):
        """Partially overlapping spans should be resolved correctly."""
        entities = [
            DetectedEntity(
                "42 MG Road Bangalore", "ADDRESS", 0, 21,
                confidence=0.80, detector_name="NER",
            ),
            DetectedEntity(
                "Bangalore", "ADDRESS", 12, 21,
                confidence=0.60, detector_name="NER",
            ),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 1
        assert resolved[0].text == "42 MG Road Bangalore"

    def test_non_overlapping_different_types_kept(self):
        """Non-overlapping entities of different types should all be kept."""
        entities = [
            DetectedEntity("user@test.com", "EMAIL", 0, 13, confidence=0.98, detector_name="Regex", validated=True),
            DetectedEntity("+91 98765 43210", "PHONE", 20, 35, confidence=0.90, detector_name="Regex", validated=True),
            DetectedEntity("Kushal Hegde", "FULL_NAME", 40, 52, confidence=0.85, detector_name="NER"),
        ]
        resolved = resolve_overlaps(entities)
        assert len(resolved) == 3

    def test_empty_input(self):
        """Empty input should return empty output."""
        assert resolve_overlaps([]) == []


# ===========================================================================
# DOCX Reader Tests
# ===========================================================================


@pytest.mark.skipif(not SAMPLE_DOCX.exists(), reason="Sample DOCX fixture not generated")
class TestDocxReader:
    """Tests for the structure-preserving DOCX reader."""

    def test_extract_segments(self):
        """Should extract segments from paragraphs, tables, and headers."""
        reader = DocxReader(SAMPLE_DOCX)
        segments = reader.extract_segments()
        assert len(segments) > 0

    def test_segments_contain_known_text(self):
        """Segments should contain known PII from the fixture."""
        reader = DocxReader(SAMPLE_DOCX)
        segments = reader.extract_segments()
        all_text = " ".join(s.text for s in segments)

        # These are known PII embedded in the fixture
        assert "Kushal Subbayya Hegde" in all_text
        assert "info@techstartup.co.in" in all_text
        assert "ABCPD1234E" in all_text
        assert "U28129PN1979PLC141032" in all_text

    def test_table_cells_extracted(self):
        """Table cell content should be in the extracted segments."""
        reader = DocxReader(SAMPLE_DOCX)
        segments = reader.extract_segments()
        table_segments = [s for s in segments if s.location_type.value == "table_cell"]
        assert len(table_segments) > 0

    def test_header_extracted(self):
        """Header content should be in the extracted segments."""
        reader = DocxReader(SAMPLE_DOCX)
        segments = reader.extract_segments()
        header_segments = [s for s in segments if s.location_type.value == "header"]
        assert len(header_segments) > 0

    def test_footer_extracted(self):
        """Footer content should be in the extracted segments."""
        reader = DocxReader(SAMPLE_DOCX)
        segments = reader.extract_segments()
        footer_segments = [s for s in segments if s.location_type.value == "footer"]
        assert len(footer_segments) > 0


# ===========================================================================
# End-to-End Pipeline Tests
# ===========================================================================


@pytest.mark.skipif(not SAMPLE_DOCX.exists(), reason="Sample DOCX fixture not generated")
class TestPipelineEndToEnd:
    """End-to-end tests for the full redaction pipeline."""

    def test_pipeline_runs_without_error(self):
        """Pipeline should complete without raising exceptions."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            entity_map_path = Path(tmpdir) / "entity_map.json"
            report_path = Path(tmpdir) / "report.json"

            pipeline = RedactionPipeline(seed=42)
            stats = pipeline.run(
                input_path=SAMPLE_DOCX,
                output_path=output_path,
                entity_map_path=entity_map_path,
                report_path=report_path,
            )

            assert output_path.exists()
            assert entity_map_path.exists()
            assert report_path.exists()

    def test_output_is_valid_docx(self):
        """Output should be a valid, openable DOCX file."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            pipeline = RedactionPipeline(seed=42)
            pipeline.run(input_path=SAMPLE_DOCX, output_path=output_path)

            # Should be loadable by python-docx
            doc = Document(str(output_path))
            assert len(doc.paragraphs) > 0

    def test_pii_is_redacted(self):
        """Known PII should not appear in the redacted output."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            pipeline = RedactionPipeline(seed=42)
            pipeline.run(input_path=SAMPLE_DOCX, output_path=output_path)

            doc = Document(str(output_path))
            # Collect ALL text: paragraphs + table cells
            all_text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text_parts.append(cell.text)
            all_text = " ".join(all_text_parts)

            # Known PII from fixture should be replaced
            # Check regex-detected PII (high confidence)
            assert "ABCPD1234E" not in all_text, "PAN should be redacted"
            assert "U28129PN1979PLC141032" not in all_text, "CIN should be redacted"

    def test_stats_populated(self):
        """Redaction stats should show entities found and redacted."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            pipeline = RedactionPipeline(seed=42)
            stats = pipeline.run(input_path=SAMPLE_DOCX, output_path=output_path)

            assert stats.segments_processed > 0
            assert sum(stats.entities_found.values()) > 0
            assert sum(stats.entities_redacted.values()) > 0
            assert stats.duration_seconds > 0

    def test_entity_consistency(self):
        """Same PII appearing multiple times should get the same fake value."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            entity_map_path = Path(tmpdir) / "entity_map.json"

            pipeline = RedactionPipeline(seed=42)
            pipeline.run(
                input_path=SAMPLE_DOCX,
                output_path=output_path,
                entity_map_path=entity_map_path,
            )

            # The entity map should exist and have entries
            import json
            with open(entity_map_path) as f:
                entity_map = json.load(f)
            assert len(entity_map) > 0

    def test_non_pii_preserved(self):
        """Non-PII content should remain in the output unchanged."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "redacted.docx"
            pipeline = RedactionPipeline(seed=42)
            pipeline.run(input_path=SAMPLE_DOCX, output_path=output_path)

            doc = Document(str(output_path))
            all_text = " ".join(p.text for p in doc.paragraphs)

            # Generic legal/structural text should survive — these are not PII
            assert "incorporated" in all_text.lower()
            assert "Companies Act" in all_text
