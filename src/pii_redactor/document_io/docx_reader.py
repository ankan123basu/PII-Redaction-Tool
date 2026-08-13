"""
Structure-preserving DOCX text extraction.

Iterates through all document elements (paragraphs, tables, headers, footers)
and returns TextSegment objects that preserve the mapping between extracted text
and the original document structure (runs, cells, etc.) so that the writer can
perform in-place replacement without losing formatting.

Design decision: We iterate elements using document.iter_inner_content() for
correct paragraph/table interleaving, plus explicit section header/footer
traversal. Each text segment carries enough metadata to map character offsets
back to individual runs, enabling run-level replacement in the writer.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from enum import StrEnum
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx.oxml.ns import qn  # noqa: F401
    from docx.section import Section

logger = logging.getLogger(__name__)


class LocationType(StrEnum):
    """Where in the document a text segment was found."""

    PARAGRAPH = "paragraph"
    TABLE_CELL = "table_cell"
    HEADER = "header"
    FOOTER = "footer"


@dataclass
class RunInfo:
    """Metadata about a single run within a text segment.

    Attributes:
        run_index: Index of this run in the parent paragraph's runs list.
        start_offset: Character offset where this run's text begins
                      within the full segment text.
        end_offset: Character offset where this run's text ends.
        text: The original text of this run.
    """

    run_index: int
    start_offset: int
    end_offset: int
    text: str


@dataclass
class TextSegment:
    """A discrete block of text extracted from the document with full provenance.

    Attributes:
        text: The concatenated text content of this segment.
        location_type: Where in the document this text was found.
        location_id: A unique identifier for locating this element in the
                     document (e.g., paragraph index, table/row/col coords,
                     section/header-or-footer).
        runs: Metadata mapping character offsets to individual runs,
              enabling run-level replacement in the writer.
        element_ref: Reference to the original python-docx paragraph object.
                     Used by the writer to modify runs in place.
    """

    text: str
    location_type: LocationType
    location_id: str
    runs: list[RunInfo] = field(default_factory=list)
    element_ref: Paragraph | None = None


class DocxReader:
    """Structure-preserving DOCX reader.

    Extracts all text from a DOCX document while preserving the mapping
    between character offsets and the underlying run structure. This allows
    the DocxWriter to perform precise, formatting-preserving replacements.

    Usage:
        reader = DocxReader("path/to/input.docx")
        segments = reader.extract_segments()
        for seg in segments:
            print(seg.text, seg.location_type, seg.location_id)
    """

    def __init__(self, file_path: str | Path) -> None:
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"Input DOCX not found: {self.file_path}")
        self.document = Document(str(self.file_path))
        logger.info("Loaded DOCX: %s", self.file_path.name)

    def extract_segments(self) -> list[TextSegment]:
        """Extract all text segments from the document.

        Returns segments from:
        1. Headers and footers (per section)
        2. Body content (paragraphs and tables in document order)

        Returns:
            List of TextSegment objects with full run metadata.
        """
        segments: list[TextSegment] = []

        # --- 1. Headers and Footers (per section) ---
        for section_idx, section in enumerate(self.document.sections):
            segments.extend(self._extract_header_footer(section, section_idx))

        # --- 2. Body content (interleaved paragraphs and tables) ---
        para_counter = 0
        table_counter = 0

        for element in self.document.iter_inner_content():
            if isinstance(element, Paragraph):
                seg = self._extract_paragraph(
                    element, LocationType.PARAGRAPH, f"body/para/{para_counter}"
                )
                if seg and seg.text.strip():
                    segments.append(seg)
                para_counter += 1
            elif isinstance(element, Table):
                segments.extend(self._extract_table(element, table_counter))
                table_counter += 1

        logger.info(
            "Extracted %d text segments (%d body paragraphs, %d tables)",
            len(segments),
            para_counter,
            table_counter,
        )
        return segments

    def _extract_header_footer(
        self, section: Section, section_idx: int
    ) -> list[TextSegment]:
        """Extract text from a section's header and footer.

        Handles the is_linked_to_previous flag: if linked, skip to avoid
        duplicating content from the previous section.
        """
        segments: list[TextSegment] = []

        # Header
        header = section.header
        if header and not header.is_linked_to_previous:
            for para_idx, para in enumerate(header.paragraphs):
                seg = self._extract_paragraph(
                    para,
                    LocationType.HEADER,
                    f"section/{section_idx}/header/para/{para_idx}",
                )
                if seg and seg.text.strip():
                    segments.append(seg)

            # Tables inside headers
            for tbl_idx, table in enumerate(header.tables):
                segments.extend(
                    self._extract_table(
                        table,
                        tbl_idx,
                        location_prefix=f"section/{section_idx}/header",
                    )
                )

        # Footer
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            for para_idx, para in enumerate(footer.paragraphs):
                seg = self._extract_paragraph(
                    para,
                    LocationType.FOOTER,
                    f"section/{section_idx}/footer/para/{para_idx}",
                )
                if seg and seg.text.strip():
                    segments.append(seg)

            # Tables inside footers
            for tbl_idx, table in enumerate(footer.tables):
                segments.extend(
                    self._extract_table(
                        table,
                        tbl_idx,
                        location_prefix=f"section/{section_idx}/footer",
                    )
                )

        return segments

    def _extract_paragraph(
        self,
        paragraph: Paragraph,
        location_type: LocationType,
        location_id: str,
    ) -> TextSegment | None:
        """Extract a single paragraph into a TextSegment with run mapping.

        Iterates runs to build a character-offset → run mapping so the writer
        can replace text at the run level without losing formatting.
        """
        runs_info: list[RunInfo] = []
        current_offset = 0

        for run_idx, run in enumerate(paragraph.runs):
            run_text = run.text or ""
            if run_text:
                runs_info.append(
                    RunInfo(
                        run_index=run_idx,
                        start_offset=current_offset,
                        end_offset=current_offset + len(run_text),
                        text=run_text,
                    )
                )
                current_offset += len(run_text)

        full_text = "".join(r.text for r in runs_info) if runs_info else paragraph.text or ""

        if not full_text:
            return None

        return TextSegment(
            text=full_text,
            location_type=location_type,
            location_id=location_id,
            runs=runs_info,
            element_ref=paragraph,
        )

    def _extract_table(
        self,
        table: Table,
        table_idx: int,
        location_prefix: str = "body",
    ) -> list[TextSegment]:
        """Extract all cells from a table, handling merged cells.

        For merged cells, python-docx returns the same cell object for
        multiple grid positions. We track seen cell IDs to avoid
        processing duplicates.

        IMPORTANT: All paragraphs within a single cell are joined into
        one TextSegment (space-separated) so that multi-line addresses
        and other multi-paragraph content reach the detector as one
        contiguous block. This prevents address fragmentation.
        """
        segments: list[TextSegment] = []
        seen_cell_ids: set[int] = set()

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Deduplicate merged cells by checking the cell's XML element identity
                cell_id = id(cell._element)
                if cell_id in seen_cell_ids:
                    continue
                seen_cell_ids.add(cell_id)

                # Join all paragraphs within this cell into one segment
                seg = self._extract_table_cell(
                    cell, table_idx, row_idx, col_idx, location_prefix
                )
                if seg and seg.text.strip():
                    segments.append(seg)

                # Handle nested tables inside cells
                for nested_tbl_idx, nested_table in enumerate(cell.tables):
                    segments.extend(
                        self._extract_table(
                            nested_table,
                            nested_tbl_idx,
                            location_prefix=(
                                f"{location_prefix}/table/{table_idx}"
                                f"/row/{row_idx}/col/{col_idx}"
                            ),
                        )
                    )

        return segments

    def _extract_table_cell(
        self,
        cell,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        location_prefix: str,
    ) -> TextSegment | None:
        """Join all paragraphs in a table cell into a single TextSegment.

        Multi-paragraph cells (common for addresses in legal tables) are
        concatenated with a space separator. The run mapping is built across
        all paragraphs so the writer can still do run-level replacement.

        The element_ref points to the first paragraph — for multi-paragraph
        cells, replacements that span paragraph boundaries fall back to the
        first paragraph's runs (the writer handles partial matches gracefully).
        """
        paragraphs = cell.paragraphs
        if not paragraphs:
            return None

        # If only one paragraph, use the standard extraction
        if len(paragraphs) == 1:
            location_id = (
                f"{location_prefix}/table/{table_idx}"
                f"/row/{row_idx}/col/{col_idx}/para/0"
            )
            return self._extract_paragraph(
                paragraphs[0], LocationType.TABLE_CELL, location_id
            )

        # Multiple paragraphs: join with space separator
        all_runs: list[RunInfo] = []
        text_parts: list[str] = []
        current_offset = 0
        first_paragraph = None

        for _para_idx, para in enumerate(paragraphs):
            para_text_parts: list[str] = []

            for run_idx, run in enumerate(para.runs):
                run_text = run.text or ""
                if run_text:
                    all_runs.append(
                        RunInfo(
                            run_index=run_idx,
                            start_offset=current_offset,
                            end_offset=current_offset + len(run_text),
                            text=run_text,
                        )
                    )
                    para_text_parts.append(run_text)
                    current_offset += len(run_text)

            para_text = "".join(para_text_parts) if para_text_parts else (para.text or "")
            if not para_text_parts and para_text:
                # Paragraph has text but no runs — track the text offset
                current_offset += len(para_text)

            if para_text.strip():
                if first_paragraph is None:
                    first_paragraph = para
                text_parts.append(para_text)
                # Add space separator between paragraphs (not after last)
                current_offset += 1  # for the space

        if not text_parts:
            return None

        full_text = " ".join(text_parts)
        location_id = (
            f"{location_prefix}/table/{table_idx}"
            f"/row/{row_idx}/col/{col_idx}/cell"
        )

        return TextSegment(
            text=full_text,
            location_type=LocationType.TABLE_CELL,
            location_id=location_id,
            runs=all_runs,
            element_ref=first_paragraph,
        )
